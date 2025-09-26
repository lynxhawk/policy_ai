#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
人才流动分析系统测试脚本
test_talent_analyzer.py
"""

import json
import pandas as pd
import os
from datetime import datetime
from talent_analyzer import TalentDataAnalyzer, AnalysisResult
from typing import Dict, List, Optional

# 尝试导入python-docx，如果没有则提示安装
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: 未安装python-docx库，将跳过Word文档生成")
    print("可通过以下命令安装: pip install python-docx")


class TalentAnalyzerTester:
    """人才分析器测试类"""
    
    def __init__(self):
        self.analyzer = TalentDataAnalyzer()
        self.test_results = {}
        
    def load_json_data(self, file_path: str) -> Optional[List[Dict]]:
        """加载JSON数据文件"""
        try:
            if not os.path.exists(file_path):
                print(f"错误: 文件 {file_path} 不存在")
                return None
                
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
            print(f"成功加载 {file_path}, 共 {len(data)} 条记录")
            return data
            
        except Exception as e:
            print(f"加载文件 {file_path} 时出错: {str(e)}")
            return None
    
    def validate_and_convert_data(self, data: List[Dict], data_type: str) -> Optional[pd.DataFrame]:
        """验证并转换数据为DataFrame"""
        try:
            if data_type in ['applicant_position', 'corporate_position']:
                # 行业数据验证
                validated_data = self.analyzer.validate_industry_data(data)
            elif data_type == 'local_count':
                # 外地户籍数据验证
                validated_data = self.analyzer.validate_nonlocal_data(data)
            else:
                raise ValueError(f"未知的数据类型: {data_type}")
                
            # 转换为DataFrame
            df = pd.DataFrame(validated_data)
            df = self.analyzer.preprocess_data(df)
            
            print(f"{data_type} 数据验证通过，转换为DataFrame成功")
            print(f"数据形状: {df.shape}")
            print(f"列名: {list(df.columns)}")
            
            return df
            
        except Exception as e:
            print(f"{data_type} 数据验证失败: {str(e)}")
            return None
    
    def test_applicant_position_analysis(self) -> Optional[List[AnalysisResult]]:
        """测试求职者求职岗位行业数据分析"""
        print("\n" + "="*50)
        print("测试求职者求职岗位行业数据分析")
        print("="*50)
        
        # 加载数据
        data = self.load_json_data('applicant_position.json')
        if data is None:
            return None
            
        # 验证和转换数据
        df = self.validate_and_convert_data(data, 'applicant_position')
        if df is None:
            return None
            
        try:
            # 执行分析
            results = self.analyzer.analyze_industry_data(df, '求职者求职岗位', 'count')
            
            print(f"分析完成，共分析 {len(results)} 个行业")
            
            # 保存结果
            self.test_results['applicant_position'] = results
            
            return results
            
        except Exception as e:
            print(f"求职者数据分析失败: {str(e)}")
            return None
    
    def test_corporate_position_analysis(self) -> Optional[List[AnalysisResult]]:
        """测试公司招聘岗位行业数据分析"""
        print("\n" + "="*50)
        print("测试公司招聘岗位行业数据分析")
        print("="*50)
        
        # 加载数据
        data = self.load_json_data('corporate_position.json')
        if data is None:
            return None
            
        # 验证和转换数据
        df = self.validate_and_convert_data(data, 'corporate_position')
        if df is None:
            return None
            
        try:
            # 执行分析
            results = self.analyzer.analyze_industry_data(df, '公司招聘岗位', 'count')
            
            print(f"分析完成，共分析 {len(results)} 个行业")
            
            # 保存结果
            self.test_results['corporate_position'] = results
            
            return results
            
        except Exception as e:
            print(f"公司招聘数据分析失败: {str(e)}")
            return None
    
    def test_local_count_analysis(self) -> Optional[AnalysisResult]:
        """测试外地户籍就业数据分析"""
        print("\n" + "="*50)
        print("测试外地户籍就业数据分析")
        print("="*50)
        
        # 加载数据
        data = self.load_json_data('local_count.json')
        if data is None:
            return None
            
        # 验证和转换数据
        df = self.validate_and_convert_data(data, 'local_count')
        if df is None:
            return None
            
        try:
            # 执行分析
            result = self.analyzer.analyze_nonlocal_data(df, 'count')
            
            if result:
                print("外地户籍数据分析完成")
                # 保存结果
                self.test_results['local_count'] = result
            else:
                print("外地户籍数据分析返回空结果")
                
            return result
            
        except Exception as e:
            print(f"外地户籍数据分析失败: {str(e)}")
            return None
    
    def format_result_for_display(self, result: AnalysisResult) -> str:
        """格式化单个分析结果用于显示"""
        lines = []
        
        if hasattr(result, 'industry') and result.industry:
            lines.append(f"行业: {result.industry}")
        lines.append(f"数据类型: {result.data_type}")
        lines.append(f"记录总数: {result.total_records}")
        lines.append(f"平均数量: {result.mean_count:.2f}")
        lines.append(f"最大值: {result.max_count:.2f}")
        lines.append(f"最小值: {result.min_count:.2f}")
        lines.append(f"总变化率: {result.total_change_pct:.4f}%")
        lines.append(f"稳定性指数: {result.stability_index:.4f} ({result.stability_rating})")
        lines.append(f"趋势强度: {result.trend_strength:.6f} ({result.trend_rating})")
        lines.append(f"波动率: {result.volatility:.4f} ({result.volatility_rating})")
        lines.append(f"平均环比增长率: {result.avg_mom_growth:.4f}%")
        lines.append(f"平均同比增长率: {result.avg_yoy_growth:.4f}%")
        lines.append(f"最新环比增长率: {result.latest_mom_growth:.4f}%")
        lines.append(f"最新同比增长率: {result.latest_yoy_growth:.4f}%")
        
        return "\n".join(lines)
    
    def generate_summary_statistics(self) -> Dict:
        """生成汇总统计信息"""
        summary = {
            'analysis_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total_datasets': len(self.test_results),
            'datasets_analyzed': []
        }
        
        for key, results in self.test_results.items():
            dataset_info = {'name': key}
            
            if isinstance(results, list):
                # 行业数据结果
                dataset_info['type'] = 'industry_analysis'
                dataset_info['industry_count'] = len(results)
                
                if results:
                    # 计算行业级别的汇总统计
                    stability_indices = [r.stability_index for r in results]
                    trend_strengths = [r.trend_strength for r in results]
                    volatilities = [r.volatility for r in results]
                    change_pcts = [r.total_change_pct for r in results]
                    
                    dataset_info['avg_stability_index'] = sum(stability_indices) / len(stability_indices)
                    dataset_info['avg_trend_strength'] = sum(trend_strengths) / len(trend_strengths)
                    dataset_info['avg_volatility'] = sum(volatilities) / len(volatilities)
                    dataset_info['avg_total_change_pct'] = sum(change_pcts) / len(change_pcts)
                    
                    # 分类统计
                    dataset_info['high_stability_count'] = len([r for r in results if r.stability_index > 80])
                    dataset_info['strong_trend_count'] = len([r for r in results if r.trend_strength > 0.05])
                    dataset_info['growing_industries'] = len([r for r in results if r.total_change_pct > 0])
                    dataset_info['declining_industries'] = len([r for r in results if r.total_change_pct < 0])
                
            elif isinstance(results, AnalysisResult):
                # 外地户籍数据结果
                dataset_info['type'] = 'nonlocal_analysis'
                dataset_info['stability_index'] = results.stability_index
                dataset_info['trend_strength'] = results.trend_strength
                dataset_info['volatility'] = results.volatility
                dataset_info['total_change_pct'] = results.total_change_pct
                dataset_info['stability_rating'] = results.stability_rating
                dataset_info['trend_rating'] = results.trend_rating
                dataset_info['volatility_rating'] = results.volatility_rating
            
            summary['datasets_analyzed'].append(dataset_info)
        
        return summary
    
    def save_results_to_txt(self, filename: str = "talent_analysis_results.txt"):
        """将结果保存为文本文件"""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write("人才流动分析系统测试结果报告\n")
                f.write("=" * 60 + "\n\n")
                
                # 写入汇总统计
                summary = self.generate_summary_statistics()
                f.write("汇总统计信息\n")
                f.write("-" * 30 + "\n")
                f.write(f"分析时间: {summary['analysis_time']}\n")
                f.write(f"分析数据集总数: {summary['total_datasets']}\n\n")
                
                # 写入详细结果
                for key, results in self.test_results.items():
                    f.write(f"\n{key.upper()} 分析结果\n")
                    f.write("=" * 60 + "\n\n")
                    
                    if isinstance(results, list):
                        # 行业数据结果
                        f.write(f"共分析 {len(results)} 个行业:\n\n")
                        
                        for i, result in enumerate(results, 1):
                            f.write(f"[{i}] ")
                            f.write(self.format_result_for_display(result))
                            f.write("\n" + "-" * 40 + "\n\n")
                            
                    elif isinstance(results, AnalysisResult):
                        # 外地户籍数据结果
                        f.write(self.format_result_for_display(results))
                        f.write("\n\n")
                
                # 写入汇总统计详情
                f.write("\n详细汇总统计\n")
                f.write("=" * 60 + "\n\n")
                
                for dataset in summary['datasets_analyzed']:
                    f.write(f"数据集: {dataset['name']}\n")
                    f.write(f"类型: {dataset['type']}\n")
                    
                    if dataset['type'] == 'industry_analysis':
                        f.write(f"行业数量: {dataset['industry_count']}\n")
                        if 'avg_stability_index' in dataset:
                            f.write(f"平均稳定性指数: {dataset['avg_stability_index']:.4f}\n")
                            f.write(f"平均趋势强度: {dataset['avg_trend_strength']:.6f}\n")
                            f.write(f"平均波动率: {dataset['avg_volatility']:.4f}\n")
                            f.write(f"平均总变化率: {dataset['avg_total_change_pct']:.4f}%\n")
                            f.write(f"高稳定性行业数量: {dataset['high_stability_count']}\n")
                            f.write(f"强趋势行业数量: {dataset['strong_trend_count']}\n")
                            f.write(f"增长行业数量: {dataset['growing_industries']}\n")
                            f.write(f"下降行业数量: {dataset['declining_industries']}\n")
                    
                    elif dataset['type'] == 'nonlocal_analysis':
                        f.write(f"稳定性指数: {dataset['stability_index']:.4f} ({dataset['stability_rating']})\n")
                        f.write(f"趋势强度: {dataset['trend_strength']:.6f} ({dataset['trend_rating']})\n")
                        f.write(f"波动率: {dataset['volatility']:.4f} ({dataset['volatility_rating']})\n")
                        f.write(f"总变化率: {dataset['total_change_pct']:.4f}%\n")
                    
                    f.write("\n" + "-" * 40 + "\n\n")
                
                f.write(f"\n报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            
            print(f"结果已保存到文件: {filename}")
            
        except Exception as e:
            print(f"保存文本文件失败: {str(e)}")
    
    def save_results_to_docx(self, filename: str = "talent_analysis_results.docx"):
        """将结果保存为Word文档"""
        if not DOCX_AVAILABLE:
            print("跳过Word文档生成（python-docx未安装）")
            return
        
        try:
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('人才流动分析系统测试结果报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加汇总信息
            summary = self.generate_summary_statistics()
            doc.add_heading('汇总统计信息', level=1)
            
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"分析时间: {summary['analysis_time']}\n")
            summary_para.add_run(f"分析数据集总数: {summary['total_datasets']}")
            
            # 添加详细结果
            for key, results in self.test_results.items():
                doc.add_heading(f'{key.upper()} 分析结果', level=1)
                
                if isinstance(results, list):
                    # 行业数据结果
                    doc.add_paragraph(f'共分析 {len(results)} 个行业:')
                    
                    for i, result in enumerate(results, 1):
                        doc.add_heading(f'行业 {i}: {result.industry if hasattr(result, "industry") else "未知"}', level=2)
                        result_text = self.format_result_for_display(result)
                        doc.add_paragraph(result_text)
                        
                elif isinstance(results, AnalysisResult):
                    # 外地户籍数据结果
                    result_text = self.format_result_for_display(results)
                    doc.add_paragraph(result_text)
            
            # 添加详细汇总统计
            doc.add_heading('详细汇总统计', level=1)
            
            for dataset in summary['datasets_analyzed']:
                doc.add_heading(f"数据集: {dataset['name']}", level=2)
                
                stats_para = doc.add_paragraph()
                stats_para.add_run(f"类型: {dataset['type']}\n")
                
                if dataset['type'] == 'industry_analysis':
                    stats_para.add_run(f"行业数量: {dataset['industry_count']}\n")
                    if 'avg_stability_index' in dataset:
                        stats_para.add_run(f"平均稳定性指数: {dataset['avg_stability_index']:.4f}\n")
                        stats_para.add_run(f"平均趋势强度: {dataset['avg_trend_strength']:.6f}\n")
                        stats_para.add_run(f"平均波动率: {dataset['avg_volatility']:.4f}\n")
                        stats_para.add_run(f"平均总变化率: {dataset['avg_total_change_pct']:.4f}%\n")
                        stats_para.add_run(f"高稳定性行业数量: {dataset['high_stability_count']}\n")
                        stats_para.add_run(f"强趋势行业数量: {dataset['strong_trend_count']}\n")
                        stats_para.add_run(f"增长行业数量: {dataset['growing_industries']}\n")
                        stats_para.add_run(f"下降行业数量: {dataset['declining_industries']}")
                
                elif dataset['type'] == 'nonlocal_analysis':
                    stats_para.add_run(f"稳定性指数: {dataset['stability_index']:.4f} ({dataset['stability_rating']})\n")
                    stats_para.add_run(f"趋势强度: {dataset['trend_strength']:.6f} ({dataset['trend_rating']})\n")
                    stats_para.add_run(f"波动率: {dataset['volatility']:.4f} ({dataset['volatility_rating']})\n")
                    stats_para.add_run(f"总变化率: {dataset['total_change_pct']:.4f}%")
            
            # 添加生成时间
            doc.add_paragraph(f"\n报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # 保存文档
            doc.save(filename)
            print(f"Word文档已保存到: {filename}")
            
        except Exception as e:
            print(f"保存Word文档失败: {str(e)}")
    
    def run_all_tests(self):
        """运行所有测试"""
        print("开始运行人才流动分析系统测试")
        print("=" * 60)
        
        # 测试求职者数据分析
        self.test_applicant_position_analysis()
        
        # 测试公司招聘数据分析
        self.test_corporate_position_analysis()
        
        # 测试外地户籍数据分析
        self.test_local_count_analysis()
        
        # 生成报告
        print("\n" + "="*60)
        print("生成测试报告")
        print("="*60)
        
        if self.test_results:
            self.save_results_to_txt()
            self.save_results_to_docx()
            print("测试完成！")
        else:
            print("没有成功的测试结果，请检查数据文件是否存在并且格式正确。")


def main():
    """主函数"""
    # 检查必要的数据文件是否存在
    required_files = ['applicant_position.json', 'corporate_position.json', 'local_count.json']
    missing_files = [f for f in required_files if not os.path.exists(f)]
    
    if missing_files:
        print("错误: 缺少以下数据文件:")
        for file in missing_files:
            print(f"  - {file}")
        print("\n请确保以下数据文件存在于当前目录:")
        print("  - applicant_position.json (求职者求职岗位行业数据)")
        print("  - corporate_position.json (公司招聘岗位行业数据)")
        print("  - local_count.json (外地户籍就业数据)")
        return
    
    # 创建测试器并运行测试
    tester = TalentAnalyzerTester()
    tester.run_all_tests()


if __name__ == "__main__":
    main()